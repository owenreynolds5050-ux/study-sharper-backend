"""
LangChain-based document processor for Study Sharper
Handles: PDF/DOCX/TXT extraction, chunking, and embedding generation
"""

import logging
import hashlib
import re
from typing import Dict, List
from pathlib import Path

import pdfplumber
from docx import Document as DocxDocument
from langchain_community.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.embeddings import HuggingFaceEmbeddings
from langchain_core.documents import Document

logger = logging.getLogger(__name__)


class LangChainProcessor:
    """Process documents: extract text, chunk, and generate embeddings."""

    def __init__(self):
        """Initialize processor with embeddings model."""
        self.embeddings = HuggingFaceEmbeddings(
            model_name="sentence-transformers/all-MiniLM-L6-v2",
            model_kwargs={"device": "cpu"},
        )
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=100,
            separators=["\n\n", "\n", " ", ""],
        )
        logger.info("LangChainProcessor initialized with all-MiniLM-L6-v2 embeddings")

    def load_document(self, file_path: str, file_type: str) -> List[Document]:
        """
        Load document based on file type using best-in-class extractors.
        - PDF: pdfplumber (preserves layout and structure)
        - DOCX: python-docx (preserves paragraphs, lists, formatting)
        - TXT: TextLoader (simple text)

        Args:
            file_path: Path to the file
            file_type: File extension (pdf, docx, txt)

        Returns:
            List of LangChain Document objects

        Raises:
            ValueError: If file type is not supported
            Exception: If file loading fails
        """
        file_type = file_type.lower().strip(".")

        try:
            if file_type == "pdf":
                return self._load_pdf_pdfplumber(file_path)
            elif file_type == "docx":
                return self._load_docx_python_docx(file_path)
            elif file_type == "txt":
                loader = TextLoader(file_path)
                docs = loader.load()
                logger.info(f"Loaded TXT file: {len(docs)} sections")
                return docs
            else:
                raise ValueError(f"Unsupported file type: {file_type}")

        except Exception as e:
            logger.error(f"Error loading {file_type} document from {file_path}: {e}")
            raise

    def _load_pdf_pdfplumber(self, file_path: str) -> List[Document]:
        """
        Load PDF using pdfplumber for better formatting preservation.
        """
        docs = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        docs.append(Document(
                            page_content=text,
                            metadata={"page": page_num, "source": file_path}
                        ))
            logger.info(f"Loaded {len(docs)} pages from PDF using pdfplumber")
            return docs
        except Exception as e:
            logger.error(f"Error loading PDF with pdfplumber: {e}")
            raise

    def _load_docx_python_docx(self, file_path: str) -> List[Document]:
        """
        Load DOCX using python-docx to preserve structure, lists, and formatting.
        """
        try:
            doc = DocxDocument(file_path)
            text_parts = []
            
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)
            
            # Also extract from tables if present
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text_parts.append(row_text)
            
            full_text = "\n".join(text_parts)
            docs = [Document(
                page_content=full_text,
                metadata={"source": file_path}
            )]
            logger.info(f"Loaded DOCX with {len(text_parts)} paragraphs/rows")
            return docs
        except Exception as e:
            logger.error(f"Error loading DOCX with python-docx: {e}")
            raise

    def extract_text(self, docs: List[Document]) -> str:
        """
        Extract, combine, and normalize text from all pages/sections.

        Args:
            docs: List of LangChain Document objects

        Returns:
            Combined and normalized text content
        """
        full_text = "\n".join([doc.page_content for doc in docs])
        # Normalize formatting for better readability
        normalized_text = self.normalize_text(full_text)
        logger.info(f"Extracted and normalized {len(normalized_text)} characters from document")
        return normalized_text

    def normalize_text(self, text: str) -> str:
        """
        Clean up and normalize extracted text:
        - Merge broken lines (lines ending without punctuation get joined with next line)
        - Preserve double newlines for paragraph breaks
        - Clean up excessive whitespace
        - Preserve bullet points and numbered lists
        - Remove PDF line break artifacts

        Args:
            text: Raw extracted text

        Returns:
            Normalized text with better formatting
        """
        # Split into lines for processing
        lines = text.split("\n")
        normalized_lines = []
        i = 0
        
        while i < len(lines):
            line = lines[i].rstrip()
            
            # Check if this is a blank line (paragraph break)
            if not line.strip():
                # Preserve paragraph breaks (but avoid excessive blank lines)
                if normalized_lines and normalized_lines[-1] != "":
                    normalized_lines.append("")
                i += 1
                continue
            
            # Check if line is a bullet point or numbered list
            is_list_item = re.match(r"^\s*([•\-\*]|\d+[.)\]]|[a-z][.)\]])\s+", line)
            
            # If line doesn't end with punctuation and isn't a list item, try to merge with next
            if (not re.search(r"[.!?:;,\-—]$", line) and 
                i + 1 < len(lines) and 
                lines[i + 1].strip() and
                not is_list_item and
                not re.match(r"^\s*([•\-\*]|\d+[.)\]]|[a-z][.)\]])\s+", lines[i + 1])):
                # Merge with next line
                merged = line + " " + lines[i + 1].strip()
                lines[i + 1] = merged
                i += 1
                continue
            
            # Clean up excessive whitespace within line
            line = re.sub(r"\s+", " ", line).strip()
            
            if line:
                normalized_lines.append(line)
            
            i += 1
        
        # Join lines and clean up excessive blank lines
        result = "\n".join(normalized_lines)
        # Replace multiple blank lines with single blank line
        result = re.sub(r"\n\n+", "\n\n", result)
        
        return result.strip()

    def chunk_text(self, text: str) -> List[Document]:
        """
        Split text into chunks with overlap.

        Args:
            text: Full text content to chunk

        Returns:
            List of LangChain Document chunks
        """
        doc = Document(page_content=text)
        chunks = self.splitter.split_documents([doc])
        logger.info(f"Split text into {len(chunks)} chunks")
        return chunks

    def generate_embeddings(self, chunks: List[Document]) -> List[List[float]]:
        """
        Generate embeddings for each chunk.

        Args:
            chunks: List of text chunks

        Returns:
            List of embedding vectors (each 384-dimensional)
        """
        embeddings_list = []
        for i, chunk in enumerate(chunks):
            embedding = self.embeddings.embed_query(chunk.page_content)
            embeddings_list.append(embedding)
            if (i + 1) % 10 == 0:
                logger.debug(f"Generated embeddings for {i + 1}/{len(chunks)} chunks")

        logger.info(f"Generated {len(embeddings_list)} embeddings")
        return embeddings_list

    def compute_content_hash(self, text: str) -> str:
        """
        Compute SHA-256 hash of text content.

        Args:
            text: Text to hash

        Returns:
            Hex-encoded SHA-256 hash
        """
        return hashlib.sha256(text.encode()).hexdigest()

    async def process_file(
        self, file_path: str, file_type: str, file_id: str, user_id: str
    ) -> Dict:
        """
        Full processing pipeline: load → extract → chunk → embed.

        Args:
            file_path: Path to uploaded file
            file_type: File type (pdf, docx, txt)
            file_id: UUID of file record
            user_id: UUID of user

        Returns:
            Dictionary with processing results:
            {
                "status": "success" | "error",
                "full_text": str,
                "chunks": List[Document],
                "embeddings": List[List[float]],
                "chunk_count": int,
                "content_hash": str,
                "error_message": str (if failed)
            }
        """
        try:
            logger.info(f"Starting file processing: file_id={file_id}, type={file_type}")

            # Step 1: Load document
            docs = self.load_document(file_path, file_type)

            # Step 2: Extract full text
            full_text = self.extract_text(docs)

            # Step 3: Split into chunks
            chunks = self.chunk_text(full_text)

            # Step 4: Generate embeddings
            embeddings = self.generate_embeddings(chunks)

            # Step 5: Compute hash
            content_hash = self.compute_content_hash(full_text)

            logger.info(
                f"File processing complete: {len(chunks)} chunks, "
                f"{len(full_text)} characters"
            )

            return {
                "status": "success",
                "full_text": full_text,
                "chunks": chunks,
                "embeddings": embeddings,
                "chunk_count": len(chunks),
                "content_hash": content_hash,
            }

        except Exception as e:
            logger.error(f"Error processing file {file_id}: {e}", exc_info=True)
            return {"status": "error", "error_message": str(e)}


# Singleton instance for use throughout the app
langchain_processor = LangChainProcessor()